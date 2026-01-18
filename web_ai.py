import streamlit as st
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document 
from io import BytesIO

# ==========================================
# 1. 核心页面配置
# ==========================================
st.set_page_config(page_title="Astra AI", page_icon="💫", layout="wide")

# 引入 CSS，彻底重定义布局
st.markdown("""
    <style>
    .stApp { background-color: #000000; color: #E0E0E0; }
    
    /* 1. 隐藏原生输入框和不必要的元素 - 仅隐藏 footer，保留 header 确保开关按钮可见 */
    div[data-testid="stChatInput"] { display: none; }
    footer { visibility: hidden; }

    /* 2. 创建 Gemini 风格的一体化底部容器 */
    .fixed-bottom-container {
        position: fixed;
        bottom: 30px;
        left: 320px; /* 避开侧边栏 */
        right: 40px;
        z-index: 999;
        background: transparent;
    }

    .gemini-capsule {
        background-color: #1E1E1E;
        border-radius: 28px;
        padding: 8px 20px;
        display: flex;
        align-items: center;
        border: 1px solid rgba(255,255,255,0.1);
        box-shadow: 0 4px 15px rgba(0,0,0,0.5);
    }

    /* 加号图标样式 */
    .plus-icon {
        color: #9E9E9E;
        font-size: 24px;
        cursor: pointer;
        margin-right: 15px;
        font-weight: 300;
        transition: color 0.2s;
    }
    .plus-icon:hover { color: #ffffff; }

    /* 文本输入区样式 */
    .custom-input {
        flex-grow: 1;
        background: transparent;
        border: none;
        color: white;
        font-size: 16px;
        outline: none;
        padding: 10px 0;
    }
    
    /* 发送按钮样式 */
    .send-btn {
        background: none;
        border: none;
        color: #757575;
        cursor: pointer;
        font-size: 20px;
        padding-left: 10px;
    }
    .send-btn:hover { color: #ffffff; }

    /* 侧边栏样式 */
    section[data-testid="stSidebar"] { background-color: #121212 !important; }
    
    .block-container { padding-bottom: 120px !important; }
    </style>
    """, unsafe_allow_html=True)

# ==========================================
# 2. 逻辑状态管理
# ==========================================
if "messages" not in st.session_state:
    st.session_state.messages = []
if "file_context" not in st.session_state:
    st.session_state.file_context = ""

# --- Word 生成辅助函数 ---
def export_to_word(msgs):
    doc = Document()
    doc.add_heading('Astra AI 对话记录', 0)
    for m in msgs:
        role = "用户" if m["role"] == "user" else "Astra AI"
        doc.add_paragraph(f"【{role}】: {m['content']}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 获取输入后的回调处理
def handle_input():
    if st.session_state.user_text:
        st.session_state.messages.append({"role": "user", "content": st.session_state.user_text})
        # 这里清除输入框
        st.session_state.user_text = ""

# API 配置
DEEPSEEK_KEY = st.secrets.get("api_key", "sk-0a477b0f3c874c8184f0a2ec168c3f2d")
client = OpenAI(api_key=DEEPSEEK_KEY, base_url="https://api.deepseek.com")

# ==========================================
# 3. 页面渲染
# ==========================================

# --- 侧边栏 ---
with st.sidebar:
    st.markdown("### 💫 Astra 历史记录")
    if st.button("➕ 开启新对话", use_container_width=True):
        st.session_state.messages = []
        st.rerun()
    
    if st.session_state.messages:
        st.markdown("---")
        st.markdown("##### 📄 文档导出")
        word_data = export_to_word(st.session_state.messages)
        st.download_button(
            label="📥 导出全部对话为 Word",
            data=word_data,
            file_name="Astra_Chat_History.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown("---")
    st.markdown("##### 📁 文档上传")
    up_file = st.file_uploader("上传 PDF 文档作为知识库", type="pdf")
    if up_file:
        reader = PdfReader(up_file)
        st.session_state.file_context = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        st.toast("文档已注入 ASTRA 核心")

# --- 主对话区 ---
if not st.session_state.messages:
    st.markdown("<div style='height: 20vh;'></div>", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align:center; color:white;'>Astra 小星AI</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center; color:#757575;'>你好！我是你的智能助手，请问有什么可以帮你的？</p>", unsafe_allow_html=True)
else:
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

# ==========================================
# 4. 【核心黑科技】一体化底部对话框
# ==========================================
st.markdown('<div class="fixed-bottom-container">', unsafe_allow_html=True)
c_icon, c_input = st.columns([0.4, 9.6])

with c_icon:
    with st.popover("＋"):
        st.write("🔧 扩展功能")
        st.toggle("开启深度联网", value=True)
        st.write("更多工具开发中...")

with c_input:
    st.text_input(
        "输入消息...", 
        key="user_text", 
        on_change=handle_input,
        label_visibility="collapsed",
        placeholder="问问 Astra，或者发送消息..."
    )
st.markdown('</div>', unsafe_allow_html=True)

# ==========================================
# 5. AI 响应逻辑
# ==========================================
if st.session_state.messages and st.session_state.messages[-1]["role"] == "user":
    last_user_msg = st.session_state.messages[-1]["content"]
    with st.chat_message("assistant"):
        with st.spinner('Astra 正在飞速思考...'):
            try:
                # --- 核心修复：在这里告诉 AI 当前的正确时间 ---
                sys_p = "你是 Astra 小星AI。今天的日期是 2026年1月18日。请专业且简洁地回答。"
                if st.session_state.file_context:
                    sys_p += f"\n背景资料: {st.session_state.file_context[:2500]}"

                res = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[{"role": "system", "content": sys_p}] + st.session_state.messages
                )
                ans = res.choices[0].message.content
                st.markdown(ans)
                st.session_state.messages.append({"role": "assistant", "content": ans})
                st.rerun()
            except Exception as e:
                st.error(f"信号微弱，请重试: {e}")
